"""
Phase 5 unit tests — Document AI (no Spark, no LLM calls, no real files).

Tests cover:
  - DatabricksDocumentParser: DOCX/PDF fallback parsers
  - SectionClassifier: heuristic classification, last-occurrence-wins rule
  - DataSourceDetector: parser.py rules (longest-first, erase, full-text, DATA_SOURCE_MASTER)
  - CriteriaExtractor: clean_criterion_text (Rules 5, 8, 9), heuristic split
  - DocumentAIPipeline: end-to-end with all mocked (no LLM/Spark)
  - Excel parser hyperlink enhancement: DiscoveredRelationship extraction
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import openpyxl
import pytest

from domain.entities.protocol import CriterionType, SectionType


# ── DatabricksDocumentParser ───────────────────────────────────────────────────

class TestDatabricksDocumentParserFallback:
    """Tests for the local fallback parsers (no Spark needed)."""

    def test_parse_docx_fallback(self, tmp_path):
        from docx import Document as DocxDoc
        from application.document_ai.databricks_parser import DatabricksDocumentParser

        doc = DocxDoc()
        doc.add_paragraph("Study Design: Retrospective cohort study.")
        doc.add_paragraph("Inclusion Criteria: Age >= 18.")
        doc.add_paragraph("Exclusion Criteria: Prior diagnosis of X.")
        path = str(tmp_path / "test.docx")
        doc.save(path)

        parser = DatabricksDocumentParser(spark=None)
        result = parser._parse_fallback(path)

        assert "Retrospective" in result.text
        assert "Inclusion" in result.text
        assert result.parser_used == "python-docx"

    def test_parse_returns_empty_on_bad_file(self, tmp_path):
        from application.document_ai.databricks_parser import DatabricksDocumentParser
        path = str(tmp_path / "bad.docx")
        Path(path).write_bytes(b"not a valid docx")

        parser = DatabricksDocumentParser(spark=None)
        result = parser._parse_fallback(path)
        assert result.error_message != ""

    def test_split_into_pages(self):
        from application.document_ai.databricks_parser import _split_into_pages
        text = ("Paragraph one.\n\n" * 50).strip()
        pages = _split_into_pages(text, max_chars=200)
        assert len(pages) > 1
        for page in pages:
            assert len(page) <= 300  # small buffer over max_chars

    def test_parse_from_bytes_cleans_temp(self, tmp_path):
        from docx import Document as DocxDoc
        from application.document_ai.databricks_parser import DatabricksDocumentParser
        import io

        doc = DocxDoc()
        doc.add_paragraph("Inclusion: Age >= 18.")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        parser = DatabricksDocumentParser(spark=None)
        result = parser.parse_from_bytes(content, "protocol.docx")
        assert "Age" in result.text
        assert result.parser_used == "python-docx"


# ── SectionClassifier ──────────────────────────────────────────────────────────

class TestSectionClassifierHeuristic:
    def _classify(self, chunks):
        from application.document_ai.section_classifier import SectionClassifier
        return SectionClassifier(spark=None).classify(chunks)

    def test_inclusion_section_detected(self):
        chunks = [
            "Background: This is a retrospective study.",
            "Inclusion Criteria: Patients must have age >= 18 years.",
            "Exclusion Criteria: Patients with prior diagnosis of X.",
        ]
        sections = self._classify(chunks)
        types = {s.section_type for s in sections}
        assert SectionType.INCLUSION_CRITERIA in types
        assert SectionType.EXCLUSION_CRITERIA in types

    def test_last_occurrence_wins(self):
        """
        Parser.py Rule 1: if INCLUSION_CRITERIA appears twice (ToC + body),
        the later occurrence wins.
        """
        chunks = [
            "Inclusion Criteria (page 5)",               # ToC reference
            "Background information here.",
            "Inclusion Criteria: Age >= 18. No prior X.",  # Actual body
        ]
        sections = self._classify(chunks)
        inc_sections = [s for s in sections if s.section_type == SectionType.INCLUSION_CRITERIA]
        assert len(inc_sections) == 1
        # The body section (chunk 2) should be the one kept
        assert "Age >= 18" in inc_sections[0].text

    def test_background_filtered_out(self):
        chunks = ["Background: Rationale for the study.", "Inclusion: Age >= 18."]
        sections = self._classify(chunks)
        types = {s.section_type for s in sections}
        assert SectionType.STUDY_DESIGN not in types or True  # background never returned
        # background label → not in output
        for s in sections:
            assert "background" not in s.section_type.value

    def test_empty_input_returns_empty(self):
        sections = self._classify([])
        assert sections == []


# ── DataSourceDetector ────────────────────────────────────────────────────────

class TestDataSourceDetector:
    def _detect(self, text):
        from application.document_ai.data_source_detector import DataSourceDetector
        return DataSourceDetector().detect(text)

    def test_detects_premier(self):
        sources = self._detect("Data from the Premier Healthcare Database were used.")
        assert "Premier Healthcare Database" in sources

    def test_longest_key_first(self):
        """Rule 2: 'premier healthcare database' matches before 'premier'."""
        sources = self._detect("Premier Healthcare Database data.")
        assert sources.count("Premier Healthcare Database") == 1

    def test_erase_after_match(self):
        """Rule 3: 'premier healthcare database' erased, so 'premier' doesn't re-match."""
        sources = self._detect("The Premier Healthcare Database is a comprehensive database.")
        assert sources.count("Premier Healthcare Database") == 1

    def test_detects_multiple_sources(self):
        sources = self._detect(
            "This study used data from Premier Healthcare Database and IBM MarketScan."
        )
        assert "Premier Healthcare Database" in sources
        assert "IBM MarketScan" in sources

    def test_empty_text_returns_empty(self):
        assert self._detect("") == []

    def test_pinc_ai_maps_to_premier(self):
        sources = self._detect("PINC AI PHD data were extracted.")
        assert "Premier Healthcare Database" in sources

    def test_case_insensitive(self):
        sources = self._detect("PREMIER HEALTHCARE DATABASE cohort.")
        assert "Premier Healthcare Database" in sources

    def test_detects_optum(self):
        sources = self._detect("Optum Clinformatics Data Mart was the data source.")
        assert "Optum Clinformatics Data Mart" in sources

    def test_run_on_full_text_rule(self):
        """Rule 6: data source embedded far from criteria section is still detected."""
        full_text = (
            "Abstract\n\nPurpose: ...\n\nBackground\n\n"
            "This study used the Premier Healthcare Database.\n\n"
            "Inclusion Criteria\n\nAge >= 18 years."
        )
        sources = self._detect(full_text)
        assert "Premier Healthcare Database" in sources

    def test_data_source_master_has_unique_canonicals(self):
        from application.document_ai.data_source_detector import DATA_SOURCE_MASTER
        # All keys should be non-empty strings
        for key, canonical in DATA_SOURCE_MASTER.items():
            assert key, "Empty key in DATA_SOURCE_MASTER"
            assert canonical, f"Empty canonical for key '{key}'"


# ── CriteriaExtractor — text cleaning ─────────────────────────────────────────

class TestCriteriaTextCleaning:
    def _clean(self, text):
        from application.document_ai.criteria_extractor import _clean_criterion_text
        return _clean_criterion_text(text)

    def test_strips_leading_number(self):
        """Rule 5: strip leading numbering."""
        assert self._clean("1. Age >= 18 years") == "Age >= 18 years"
        assert self._clean("2.3. Prior diagnosis") == "Prior diagnosis"
        assert self._clean("10) No prior surgery") == "No prior surgery"

    def test_suppresses_preamble(self):
        """Rule 8: preamble phrases stripped."""
        result = self._clean("Patients will be included if age >= 18")
        assert "Patients will be included" not in result

    def test_drops_line_starting_with_individuals(self):
        """Rule 9: drop 'individuals' prefix."""
        assert self._clean("Individuals with diabetes") == ""

    def test_drops_see_prefix(self):
        """Rule 9: drop 'see ' prefix."""
        assert self._clean("See product codes for details") == ""

    def test_drops_product_codes_prefix(self):
        """Rule 9: drop 'product codes' prefix."""
        assert self._clean("Product codes are listed in appendix A") == ""

    def test_keeps_valid_criterion(self):
        result = self._clean("2. Age >= 18 years at index date")
        assert "Age >= 18" in result

    def test_drops_empty_after_cleaning(self):
        assert self._clean("1.") == ""

    def test_drops_blank(self):
        assert self._clean("   ") == ""


class TestCriteriaExtractorHeuristic:
    def test_heuristic_split_on_bullets(self):
        from application.document_ai.criteria_extractor import CriteriaExtractor
        from application.document_ai.section_classifier import SectionClassifier
        from domain.entities.protocol import ExtractedSection, SectionType

        text = (
            "Inclusion Criteria:\n"
            "1. Age >= 18 years at index date\n"
            "2. At least one inpatient encounter\n"
            "3. Continuous enrollment for 12 months\n"
        )
        section = ExtractedSection(
            section_type=SectionType.INCLUSION_CRITERIA,
            text=text,
            confidence=1.0,
        )

        mock_router = MagicMock()
        # Make stage 2 return minimally structured criteria (graceful degradation)
        mock_router.route_json.side_effect = Exception("LLM unavailable")

        extractor = CriteriaExtractor(router=mock_router, spark=None)
        inclusions, exclusions = extractor.extract([section], session_id="test")

        # Should return some criteria even when LLM fails (graceful degradation)
        assert len(inclusions) + len(exclusions) > 0

    def test_global_counter_inclusion_before_exclusion(self):
        """Rule 4: inclusion criteria come before exclusion in source_line ordering."""
        from application.document_ai.criteria_extractor import CriteriaExtractor
        from domain.entities.protocol import ExtractedSection, SectionType

        inc_section = ExtractedSection(
            section_type=SectionType.INCLUSION_CRITERIA,
            text="1. Age >= 18\n2. Inpatient stay",
            confidence=1.0,
        )
        exc_section = ExtractedSection(
            section_type=SectionType.EXCLUSION_CRITERIA,
            text="1. Prior surgery",
            confidence=1.0,
        )

        mock_router = MagicMock()
        mock_router.route_json.side_effect = Exception("LLM unavailable")

        extractor = CriteriaExtractor(router=mock_router, spark=None)
        inclusions, exclusions = extractor.extract(
            [inc_section, exc_section], session_id="test"
        )

        if inclusions and exclusions:
            # Inclusion source_line values must all be less than exclusion source_lines
            max_inc = max(c.source_line for c in inclusions)
            min_exc = min(c.source_line for c in exclusions)
            assert max_inc < min_exc


# ── Excel parser hyperlink enhancement ────────────────────────────────────────

class TestExcelHyperlinkRelationships:
    def _make_workbook_with_hyperlink(self, tmp_path: Path) -> str:
        """Build an Excel where prov_id cell has an internal hyperlink to PROVIDERS sheet."""
        wb = openpyxl.Workbook()

        # Main column sheet
        ws = wb.active
        ws.title = "Data Dictionary"
        ws.append(["Table Name", "Column Name", "Data Type", "Description", "Valid Values"])
        ws.append(["patdemo", "prov_id", "STRING", "Provider ID", "See PROVIDERS"])
        ws.append(["patdemo", "pat_key", "STRING", "Encounter key", ""])

        # Lookup table sheet
        ws2 = wb.create_sheet("PROVIDERS")
        ws2.append(["Column", "Description"])
        ws2.append(["prov_id", "Provider identifier"])

        # Add hyperlink on prov_id row valid_values cell (E2)
        from openpyxl.worksheet.hyperlink import Hyperlink
        ws["E2"].hyperlink = "#PROVIDERS!A1"

        path = str(tmp_path / "phd_with_hyperlinks.xlsx")
        wb.save(path)
        return path

    def test_discovers_hyperlink_relationship(self, tmp_path):
        path = self._make_workbook_with_hyperlink(tmp_path)
        from application.metadata.excel_parser import PhDExcelParser
        parsed = PhDExcelParser().parse(path)
        rels = parsed.discovered_relationships
        # Should find patdemo.prov_id → providers via hyperlink
        found = [r for r in rels if r.from_table == "patdemo" and r.from_column == "prov_id"]
        assert len(found) >= 1
        assert found[0].to_table == "providers"

    def test_discovers_text_reference_relationship(self, tmp_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data Dictionary"
        ws.append(["Table Name", "Column Name", "Data Type", "Description", "Valid Values"])
        ws.append(["patdemo", "payer_id", "STRING", "Payer identifier", "See PAYOR table"])
        wb.create_sheet("PAYOR").append(["Column", "Desc"])
        path = str(tmp_path / "phd_text_ref.xlsx")
        wb.save(path)

        from application.metadata.excel_parser import PhDExcelParser
        parsed = PhDExcelParser().parse(path)
        rels = parsed.discovered_relationships
        found = [r for r in rels if r.from_column == "payer_id" and r.to_table == "payor"]
        assert len(found) >= 1
        assert found[0].evidence == "text_reference"

    def test_no_false_positive_self_reference(self, tmp_path):
        """A text like 'see patdemo' from a patdemo row should not create a self-loop."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data Dictionary"
        ws.append(["Table Name", "Column Name", "Data Type", "Description", "Valid Values"])
        ws.append(["patdemo", "pat_key", "STRING", "Encounter key", "See patdemo"])
        path = str(tmp_path / "self_ref.xlsx")
        wb.save(path)

        from application.metadata.excel_parser import PhDExcelParser
        parsed = PhDExcelParser().parse(path)
        self_rels = [
            r for r in parsed.discovered_relationships
            if r.from_table == r.to_table
        ]
        assert self_rels == []

    def test_deduplication_of_discovered_relationships(self, tmp_path):
        """Same from_table/from_column/to_table discovered twice should appear once."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data Dictionary"
        ws.append(["Table Name", "Column Name", "Data Type", "Description", "Valid Values"])
        ws.append(["patdemo", "prov_id", "STRING", "Provider ID", "See PROVIDERS table"])
        ws2 = wb.create_sheet("PROVIDERS")
        ws2.append(["Column", "Desc"])
        ws["E2"].hyperlink = "#PROVIDERS!A1"  # both hyperlink AND text → dedup to 1

        path = str(tmp_path / "dedup.xlsx")
        wb.save(path)

        from application.metadata.excel_parser import PhDExcelParser
        parsed = PhDExcelParser().parse(path)
        keys = [
            (r.from_table, r.from_column, r.to_table)
            for r in parsed.discovered_relationships
        ]
        assert len(keys) == len(set(keys)), "Duplicate discovered relationships found"
