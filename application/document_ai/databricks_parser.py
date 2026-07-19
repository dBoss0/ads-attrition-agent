"""
Databricks Document Parser — uses ai_parse_document() to extract text.

Requires: Databricks Runtime 14.3 LTS or later, Foundation Model API enabled.
Reference: https://docs.databricks.com/en/sql/language-manual/functions/ai_parse_document.html

ai_parse_document(content BINARY) → STRUCT<text STRING, error_message STRING>

The function is called via Spark SQL on a temp view that wraps the binary file
content. For files already in a Unity Catalog Volume, we pass the volume path
to read_files() and feed the binary content directly.

Fallback for local dev / testing (no Databricks): uses python-docx + pdfplumber.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from pyspark.sql import SparkSession

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Raw output from ai_parse_document or the fallback parsers."""
    text: str
    pages: list[str] = field(default_factory=list)  # text per page/section
    source_path: str = ""
    parser_used: str = ""
    error_message: str = ""


class DatabricksDocumentParser:
    """
    Parses protocol documents (PDF, DOCX) using Databricks ai_parse_document().

    On Databricks (SparkSession available):
        → ai_parse_document() via Spark SQL  (requires DBR 14.3+, FMA enabled)

    On local dev (no Spark / Databricks env):
        → python-docx for DOCX, pdfplumber for PDF

    The caller does NOT need to branch — this class handles detection.
    """

    def __init__(self, spark: "SparkSession | None" = None) -> None:
        self._spark = spark

    def parse_from_volume(self, volume_path: str) -> ParsedDocument:
        """
        Parse a document already uploaded to a Unity Catalog Volume.
        volume_path: e.g. /Volumes/ads_automation/main/protocols/MyProtocol.pdf
        """
        if self._spark and self._is_databricks():
            return self._parse_with_ai_parse_document(volume_path)
        return self._parse_fallback(volume_path)

    def parse_from_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse raw bytes (from Gradio upload or Volume read).
        Writes to a temp file, then delegates to the appropriate parser.
        """
        import tempfile, os
        suffix = Path(filename).suffix.lower() or ".pdf"
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        try:
            tmp.write(content)
            tmp.close()
            return self._parse_fallback(tmp.name, source_label=filename)
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

    # ── Databricks ai_parse_document ──────────────────────────────────────────

    def _parse_with_ai_parse_document(self, volume_path: str) -> ParsedDocument:
        """
        Uses Databricks built-in ai_parse_document() SQL function.

        Calls read_files() to read binary content from the Volume, then passes
        the BINARY content to ai_parse_document().

        ai_parse_document returns STRUCT<text STRING, error_message STRING>.
        DBR 14.3+, Foundation Model API must be enabled in workspace settings.
        """
        logger.info("Using ai_parse_document for: %s", volume_path)
        try:
            rows = self._spark.sql(f"""
                SELECT
                    ai_parse_document(content).text          AS doc_text,
                    ai_parse_document(content).error_message AS doc_error
                FROM read_files(
                    '{volume_path}',
                    format          => 'binaryFile',
                    pathGlobFilter  => '*'
                )
            """).collect()

            if not rows:
                return ParsedDocument(
                    text="",
                    source_path=volume_path,
                    parser_used="ai_parse_document",
                    error_message="No content returned from ai_parse_document",
                )

            row = rows[0]
            error = row["doc_error"] or ""
            text = row["doc_text"] or ""

            if error:
                logger.warning("ai_parse_document error for %s: %s", volume_path, error)

            return ParsedDocument(
                text=text,
                pages=_split_into_pages(text),
                source_path=volume_path,
                parser_used="ai_parse_document",
                error_message=error,
            )

        except Exception as exc:
            logger.warning(
                "ai_parse_document failed (%s) — falling back to python parser", exc
            )
            return self._parse_fallback(volume_path)

    # ── Local / fallback parsers ───────────────────────────────────────────────

    def _parse_fallback(self, path: str, source_label: str = "") -> ParsedDocument:
        ext = Path(path).suffix.lower()
        label = source_label or path
        if ext == ".docx":
            return self._parse_docx(path, label)
        if ext == ".pdf":
            return self._parse_pdf(path, label)
        # Unknown format — try DOCX, then PDF
        try:
            return self._parse_docx(path, label)
        except Exception:
            return self._parse_pdf(path, label)

    @staticmethod
    def _parse_docx(path: str, label: str) -> ParsedDocument:
        try:
            from docx import Document
            doc = Document(path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            logger.info("python-docx parsed %d paragraphs from %s", len(paragraphs), label)
            return ParsedDocument(
                text=full_text,
                pages=paragraphs,
                source_path=label,
                parser_used="python-docx",
            )
        except Exception as exc:
            return ParsedDocument(
                text="",
                source_path=label,
                parser_used="python-docx",
                error_message=str(exc),
            )

    @staticmethod
    def _parse_pdf(path: str, label: str) -> ParsedDocument:
        try:
            import pdfplumber
            pages_text: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text() or ""
                    pages_text.append(extracted.strip())
            full_text = "\n\n".join(p for p in pages_text if p)
            logger.info("pdfplumber parsed %d pages from %s", len(pages_text), label)
            return ParsedDocument(
                text=full_text,
                pages=pages_text,
                source_path=label,
                parser_used="pdfplumber",
            )
        except Exception as exc:
            return ParsedDocument(
                text="",
                source_path=label,
                parser_used="pdfplumber",
                error_message=str(exc),
            )

    @staticmethod
    def _is_databricks() -> bool:
        """True when running inside a Databricks cluster."""
        try:
            import subprocess
            result = subprocess.run(
                ["bash", "-c", "echo $DATABRICKS_RUNTIME_VERSION"],
                capture_output=True, text=True, timeout=2,
            )
            return bool(result.stdout.strip())
        except Exception:
            return False


# ── Helpers ────────────────────────────────────────────────────────────────────

def _split_into_pages(text: str, max_chars: int = 3000) -> list[str]:
    """Split long text into page-like chunks for downstream processing."""
    if not text:
        return []
    chunks: list[str] = []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        if current_len + len(para) > max_chars and current:
            chunks.append("\n\n".join(current))
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += len(para)
    if current:
        chunks.append("\n\n".join(current))
    return chunks
